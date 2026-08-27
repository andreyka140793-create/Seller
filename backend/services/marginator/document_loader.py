"""
Загрузка прайс-листов.

Поддерживаемые форматы (по рынку поставщиков / 1С / площадок):
  Таблицы:  .xlsx .xls .xlsm .xlsb .ods .csv .tsv .txt .dat
  Документы:.docx .pdf .html .htm
  Каталоги: .yml .xml (YML / CommerceML / произвольный XML)
  Данные:   .json .dbf
  Картинки: .jpg .jpeg .png .webp .bmp .gif
  Архивы:   .zip (внутри — любой из поддерживаемых файлов)

Не берём без доп. инфраструктуры: .rar, .mdb/.accdb, EDI binary, Google Sheets API.
"""
from __future__ import annotations

import io
import json
import logging
import zipfile
from dataclasses import dataclass
from xml.etree import ElementTree as ET

import pandas as pd

from services.marginator.file_io import read_table

logger = logging.getLogger(__name__)

TABLE_EXTENSIONS = {
    ".xlsx", ".xls", ".xlsm", ".xlsb", ".ods",
    ".csv", ".tsv", ".txt", ".text", ".md", ".dat",
}
DOC_EXTENSIONS = {".docx", ".doc"}
PDF_EXTENSIONS = {".pdf"}
HTML_EXTENSIONS = {".html", ".htm"}
XML_EXTENSIONS = {".xml", ".yml", ".yaml"}  # yml прайсы часто XML
JSON_EXTENSIONS = {".json"}
DBF_EXTENSIONS = {".dbf"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif"}
ARCHIVE_EXTENSIONS = {".zip"}

ALL_SUPPORTED = (
    TABLE_EXTENSIONS
    | DOC_EXTENSIONS
    | PDF_EXTENSIONS
    | HTML_EXTENSIONS
    | XML_EXTENSIONS
    | JSON_EXTENSIONS
    | DBF_EXTENSIONS
    | IMAGE_EXTENSIONS
    | ARCHIVE_EXTENSIONS
)


@dataclass
class LoadedDocument:
    kind: str  # table | text | image
    file_name: str
    dataframe: pd.DataFrame | None = None
    text: str | None = None
    image_bytes: bytes | None = None
    image_mime: str | None = None


def extension_of(file_name: str) -> str:
    name = (file_name or "").lower().strip()
    if "." not in name:
        return ""
    return "." + name.rsplit(".", 1)[-1]


def friendly_load_error(exc: BaseException, file_name: str = "") -> str:
    msg = str(exc) or type(exc).__name__
    low = msg.lower()
    name = file_name or "файл"
    if "password" in low or "encrypted" in low:
        return f"«{name}» защищён паролем. Снимите защиту и пришлите снова."
    if "xlrd" in low:
        return f"Старый .xls «{name}» не прочитался. Сохраните как .xlsx."
    if "xlsb" in low or "pyxlsb" in low:
        return f".xlsb «{name}»: сохраните как .xlsx."
    if "odf" in low or "ods" in low:
        return f".ods «{name}»: сохраните как .xlsx."
    if "dbf" in low:
        return f".dbf «{name}»: экспортируйте в CSV или Excel."
    if "empty" in low or "no columns" in low:
        return f"«{name}» пустой или без таблицы."
    if "zip" in low and "поддержива" in low:
        return f"В ZIP нет прайса (xlsx/csv/xml…)."
    if "pdf" in low:
        return f"PDF без текста: нужен Excel или скриншот."
    return f"Не удалось прочитать «{name}»: {msg[:220]}"


def is_supported(file_name: str) -> bool:

    return extension_of(file_name) in ALL_SUPPORTED


def supported_formats_hint() -> str:
    return (
        "xlsx, xls, xlsm, xlsb, ods, csv, tsv, txt, dat, "
        "docx, pdf, html, yml, xml, json, dbf, "
        "jpg/png/webp, zip"
    )


def load_document(file_bytes: bytes, file_name: str) -> LoadedDocument:
    ext = extension_of(file_name)
    name = file_name or "file"

    if ext in ARCHIVE_EXTENSIONS:
        return _load_from_zip(file_bytes, name)

    if ext in {".xlsx", ".xls", ".xlsm", ".xlsb"}:
        df = _read_excel_safe(file_bytes, name)
        return LoadedDocument(kind="table", file_name=name, dataframe=df)

    if ext == ".ods":
        df = _read_ods(file_bytes)
        return LoadedDocument(kind="table", file_name=name, dataframe=df)

    if ext in {".csv", ".dat"}:
        df = _read_csv_flexible(file_bytes)
        return LoadedDocument(kind="table", file_name=name, dataframe=df)

    if ext == ".tsv":
        df = pd.read_csv(io.BytesIO(file_bytes), sep="\t", header=None, dtype=str, engine="python")
        return LoadedDocument(kind="table", file_name=name, dataframe=df)

    if ext in {".txt", ".text", ".md"}:
        text = _decode_text(file_bytes)
        df = _try_text_as_table(text)
        if df is not None and len(df.columns) >= 2:
            return LoadedDocument(kind="table", file_name=name, dataframe=df, text=text)
        return LoadedDocument(kind="text", file_name=name, text=text)

    if ext == ".docx":
        df_tab = _read_docx_table(file_bytes)
        text = _read_docx(file_bytes)
        if df_tab is not None and df_tab.shape[1] >= 2 and len(df_tab) >= 1:
            return LoadedDocument(kind="table", file_name=name, dataframe=df_tab, text=text)
        df = _try_text_as_table(text)
        if df is not None and len(df.columns) >= 2:
            return LoadedDocument(kind="table", file_name=name, dataframe=df, text=text)
        return LoadedDocument(kind="text", file_name=name, text=text)

    if ext == ".doc":
        raise RuntimeError(
            "Старый .doc не поддерживается. Сохраните как .docx или .xlsx."
        )

    if ext in PDF_EXTENSIONS:
        text = _read_pdf(file_bytes)
        df = _try_text_as_table(text)
        if df is not None and len(df.columns) >= 2:
            return LoadedDocument(kind="table", file_name=name, dataframe=df, text=text)
        return LoadedDocument(kind="text", file_name=name, text=text)

    if ext in HTML_EXTENSIONS:
        text = _read_html_text(file_bytes)
        df = _try_html_tables(file_bytes)
        if df is not None and not df.empty:
            return LoadedDocument(kind="table", file_name=name, dataframe=df, text=text)
        df2 = _try_text_as_table(text)
        if df2 is not None and len(df2.columns) >= 2:
            return LoadedDocument(kind="table", file_name=name, dataframe=df2, text=text)
        return LoadedDocument(kind="text", file_name=name, text=text)

    if ext in XML_EXTENSIONS:
        df = _read_xml_price(file_bytes)
        return LoadedDocument(kind="table", file_name=name, dataframe=df)

    if ext in JSON_EXTENSIONS:
        df = _read_json_price(file_bytes)
        return LoadedDocument(kind="table", file_name=name, dataframe=df)

    if ext in DBF_EXTENSIONS:
        df = _read_dbf(file_bytes)
        return LoadedDocument(kind="table", file_name=name, dataframe=df)

    if ext in IMAGE_EXTENSIONS:
        mime = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".webp": "image/webp",
            ".bmp": "image/bmp",
            ".gif": "image/gif",
        }.get(ext, "image/jpeg")
        return LoadedDocument(
            kind="image",
            file_name=name,
            image_bytes=file_bytes,
            image_mime=mime,
        )

    raise RuntimeError(f"Формат {ext or '(без расширения)'} не поддерживается. {supported_formats_hint()}")


# ── helpers ──────────────────────────────────────────────────────────


def _decode_text(file_bytes: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _read_excel_safe(file_bytes: bytes, file_name: str) -> pd.DataFrame:
    try:
        return read_table(file_bytes, file_name, header=None)
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"Ошибка чтения Excel «{file_name}»: {e}") from e


def _read_ods(file_bytes: bytes) -> pd.DataFrame:
    try:
        return pd.read_excel(io.BytesIO(file_bytes), engine="odf", header=None, dtype=str)
    except ImportError as e:
        raise RuntimeError("Для .ods установите odfpy: pip install odfpy") from e
    except Exception as e:
        raise RuntimeError(f"Ошибка чтения ODS: {e}") from e


def _read_csv_flexible(file_bytes: bytes) -> pd.DataFrame:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        for sep in (";", ",", "\t", "|"):
            try:
                df = pd.read_csv(
                    io.BytesIO(file_bytes),
                    sep=sep,
                    header=None,
                    dtype=str,
                    engine="python",
                    encoding=enc,
                )
                if df.shape[1] >= 2:
                    return df
            except Exception:
                continue
    return pd.read_csv(io.BytesIO(file_bytes), header=None, dtype=str, engine="python")


def _try_text_as_table(text: str) -> pd.DataFrame | None:
    if not text or not text.strip():
        return None
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if len(lines) < 2:
        return None
    for sep in ("\t", ";", ",", "|"):
        rows = [ln.split(sep) for ln in lines[:50]]
        widths = {len(r) for r in rows}
        if len(widths) == 1 and list(widths)[0] >= 2:
            try:
                return pd.read_csv(
                    io.StringIO("\n".join(lines)),
                    sep=sep,
                    header=None,
                    dtype=str,
                    engine="python",
                )
            except Exception:
                continue
    return None


def _read_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("Установите python-docx для .docx") from e
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    return "\n".join(parts)


def _read_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("Установите pypdf для PDF") from e
    reader = PdfReader(io.BytesIO(file_bytes))
    parts: list[str] = []
    for page in reader.pages[:40]:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t)
    if not parts:
        raise RuntimeError(
            "В PDF нет текстового слоя. Пришлите скриншот (JPG/PNG) или Excel."
        )
    return "\n".join(parts)


def _read_html_text(file_bytes: bytes) -> str:
    text = _decode_text(file_bytes)
    # грубо убрать теги
    import re
    return re.sub(r"<[^>]+>", " ", text)


def _try_html_tables(file_bytes: bytes) -> pd.DataFrame | None:
    try:
        tables = pd.read_html(io.BytesIO(file_bytes))
        if not tables:
            return None
        # берём самую «широкую» таблицу
        tables = sorted(tables, key=lambda d: d.shape[1] * d.shape[0], reverse=True)
        df = tables[0]
        df = df.astype(str)
        df.columns = [str(c) for c in df.columns]
        return df
    except Exception:
        return None


def _local(tag: str) -> str:
    if "}" in tag:
        return tag.rsplit("}", 1)[-1]
    return tag


def _read_xml_price(file_bytes: bytes) -> pd.DataFrame:
    """YML / CommerceML / плоский XML → таблица name, price, article, …"""
    try:
        root = ET.fromstring(file_bytes)
    except ET.ParseError as e:
        # иногда yml с BOM
        text = _decode_text(file_bytes)
        try:
            root = ET.fromstring(text.encode("utf-8"))
        except Exception as e2:
            raise RuntimeError(f"Не удалось разобрать XML/YML: {e2}") from e

    rows: list[dict] = []

    # YML: shop/offers/offer
    for el in root.iter():
        if _local(el.tag).lower() != "offer":
            continue
        row: dict = {}
        if el.get("id"):
            row["Артикул"] = el.get("id")
        for child in list(el):
            key = _local(child.tag)
            val = (child.text or "").strip()
            if not val and child.get("value"):
                val = child.get("value") or ""
            if not val:
                continue
            kl = key.lower()
            if kl in ("name", "model", "title", "наименование", "название"):
                row["Наименование"] = val
            elif kl in ("price", "цена", "purchase_price", "закуп"):
                row["Цена"] = val
            elif kl in ("vendorcode", "vendor_code", "article", "артикул", "sku", "code"):
                row["Артикул"] = row.get("Артикул") or val
            elif kl in ("vendor", "brand", "производитель"):
                row["Бренд"] = val
            elif kl in ("currencyid", "currency"):
                row["Валюта"] = val
            elif kl in ("count", "quantity", "остаток", "stock"):
                row["Остаток"] = val
            else:
                row[key] = val
        if row.get("Наименование") or row.get("Цена"):
            rows.append(row)

    # CommerceML: Предложение / Товар
    if not rows:
        for el in root.iter():
            loc = _local(el.tag)
            if loc not in ("Предложение", "Товар", "Offer", "Product"):
                continue
            row = {}
            for child in list(el):
                key = _local(child.tag)
                val = (child.text or "").strip()
                if not val:
                    continue
                if key in ("Наименование", "Значение", "Name"):
                    row["Наименование"] = row.get("Наименование") or val
                elif key in ("Артикул", "Ид", "Id"):
                    row["Артикул"] = row.get("Артикул") or val
                elif key in ("ЦенаЗаЕдиницу", "Цена", "Price"):
                    row["Цена"] = val
                elif key == "Цены":
                    for p in child.iter():
                        if _local(p.tag) in ("ЦенаЗаЕдиницу", "Цена") and (p.text or "").strip():
                            row["Цена"] = (p.text or "").strip()
                            break
                else:
                    row[key] = val
            # цена может быть вложенной
            if "Цена" not in row:
                for p in el.iter():
                    if _local(p.tag) in ("ЦенаЗаЕдиницу", "Цена") and (p.text or "").strip():
                        row["Цена"] = (p.text or "").strip()
                        break
            if row.get("Наименование") or row.get("Цена"):
                rows.append(row)

    if not rows:
        # fallback: все элементы с текстом → две колонки tag/text
        flat = []
        for el in root.iter():
            if el.text and el.text.strip() and len(list(el)) == 0:
                flat.append({"Поле": _local(el.tag), "Значение": el.text.strip()})
        if len(flat) >= 2:
            return pd.DataFrame(flat)
        raise RuntimeError(
            "В XML/YML не найдены товары (offer / Предложение). "
            "Пришлите Excel или CSV."
        )

    return pd.DataFrame(rows)


def _read_json_price(file_bytes: bytes) -> pd.DataFrame:
    text = _decode_text(file_bytes)
    data = json.loads(text)

    if isinstance(data, list):
        if not data:
            raise RuntimeError("JSON-массив пуст")
        if all(isinstance(x, dict) for x in data):
            return pd.json_normalize(data)
        return pd.DataFrame({"value": data})

    if isinstance(data, dict):
        for key in ("offers", "items", "products", "goods", "data", "rows", "Предложения"):
            val = data.get(key)
            if isinstance(val, list) and val and isinstance(val[0], dict):
                return pd.json_normalize(val)
        # одна запись-товар
        if any(k in data for k in ("name", "price", "title", "Наименование", "Цена")):
            return pd.json_normalize([data])
        # dict of dicts
        if data and all(isinstance(v, dict) for v in data.values()):
            rows = []
            for k, v in data.items():
                row = dict(v)
                row.setdefault("id", k)
                rows.append(row)
            return pd.DataFrame(rows)
        return pd.json_normalize(data)

    raise RuntimeError("Неподдерживаемая структура JSON для прайса")


def _read_dbf(file_bytes: bytes) -> pd.DataFrame:
    try:
        from dbfread import DBF
    except ImportError as e:
        raise RuntimeError("Для .dbf установите dbfread: pip install dbfread") from e
    # dbfread хочет путь или file-like с именем — используем временный буфер через path-like
    import tempfile
    import os
    fd, path = tempfile.mkstemp(suffix=".dbf")
    try:
        os.write(fd, file_bytes)
        os.close(fd)
        table = DBF(path, encoding="cp1251", ignore_missing_memofile=True)
        rows = [dict(r) for r in table]
        if not rows:
            table = DBF(path, encoding="utf-8", ignore_missing_memofile=True)
            rows = [dict(r) for r in table]
        return pd.DataFrame(rows)
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


def _load_from_zip(file_bytes: bytes, archive_name: str) -> LoadedDocument:
    try:
        zf = zipfile.ZipFile(io.BytesIO(file_bytes))
    except zipfile.BadZipFile as e:
        raise RuntimeError("Повреждённый ZIP-архив") from e

    candidates = []
    for info in zf.infolist():
        if info.is_dir():
            continue
        base = info.filename.split("/")[-1]
        if base.startswith("."):
            continue
        if is_supported(base) and extension_of(base) not in ARCHIVE_EXTENSIONS:
            candidates.append(info)

    # приоритет: таблицы > xml/json > docs > images
    def prio(info):
        ext = extension_of(info.filename)
        if ext in {".xlsx", ".xls", ".xlsm", ".csv", ".tsv"}:
            return 0
        if ext in XML_EXTENSIONS | JSON_EXTENSIONS | {".ods", ".dbf"}:
            return 1
        if ext in DOC_EXTENSIONS | PDF_EXTENSIONS | HTML_EXTENSIONS:
            return 2
        return 3

    candidates.sort(key=lambda i: (prio(i), -i.file_size))
    if not candidates:
        raise RuntimeError(
            "В ZIP нет поддерживаемых прайсов. "
            f"Ожидаются: {supported_formats_hint()}"
        )

    chosen = candidates[0]
    inner_name = chosen.filename.split("/")[-1]
    inner_bytes = zf.read(chosen)
    doc = load_document(inner_bytes, inner_name)
    doc.file_name = f"{archive_name}:{inner_name}"
    return doc
